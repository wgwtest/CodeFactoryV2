from __future__ import annotations

from functools import lru_cache
from typing import Iterator

from app.parsing.models import ParsedDocument, ParsedSegment
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def parse_docx(file_path: str) -> list[str]:
    return [segment.content for segment in parse_docx_segments(file_path).segments]


def parse_docx_segments(file_path: str, *, formal_extraction_mode: bool = False) -> ParsedDocument:
    converter = _build_docling_converter()
    if converter is not None:
        parsed = _parse_docx_with_docling(file_path, converter)
        if parsed.segments:
            return parsed
    if formal_extraction_mode:
        raise ValueError(f"正式知识库抽取要求 DOC/DOCX 使用 Docling 解析，但当前文件未能通过 Docling 成功解析：{file_path}")

    parsed = _parse_docx_with_unstructured(file_path)
    if parsed.segments:
        return parsed
    if formal_extraction_mode:
        raise ValueError(f"正式知识库抽取禁止 DOC/DOCX 解析降级到非 Docling 解析器：{file_path}")

    return _parse_docx_with_python_docx(file_path)


@lru_cache(maxsize=1)
def _build_docling_converter():
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        return None

    return DocumentConverter()


def _parse_docx_with_docling(file_path: str, converter) -> ParsedDocument:
    try:
        result = converter.convert(file_path)
    except Exception:
        return ParsedDocument(parser_name="docling_docx", parser_version="v3", segments=[])

    markdown = result.document.export_to_markdown()
    return ParsedDocument(
        parser_name="docling_docx",
        parser_version="v3",
        segments=_segments_from_text(markdown, parser_name="docling_docx"),
    )


def _parse_docx_with_unstructured(file_path: str) -> ParsedDocument:
    try:
        from unstructured.partition.docx import partition_docx
    except ImportError:
        return ParsedDocument(parser_name="unstructured_docx", parser_version="v3", segments=[])

    try:
        elements = partition_docx(filename=file_path)
    except Exception:
        return ParsedDocument(parser_name="unstructured_docx", parser_version="v3", segments=[])

    segments: list[ParsedSegment] = []
    for index, element in enumerate(elements, start=1):
        text = getattr(element, "text", "").strip()
        if not text:
            continue
        metadata = getattr(element, "metadata", None)
        page_number = getattr(metadata, "page_number", None) if metadata else None
        category = getattr(element, "category", element.__class__.__name__).lower()
        heading = text.splitlines()[0][:255]
        anchor: dict[str, str | int] = {"section": heading, "line_start": index, "line_end": index}
        if page_number is not None:
            anchor["page"] = int(page_number)
        segments.append(
            ParsedSegment(
                block_id=f"docx-unstructured-{index}",
                heading=heading,
                content=text,
                anchor=anchor,
                block_type=category or "paragraph",
            )
        )

    return ParsedDocument(parser_name="unstructured_docx", parser_version="v3", segments=segments)


def _parse_docx_with_python_docx(file_path: str) -> ParsedDocument:
    document = Document(file_path)
    segments: list[ParsedSegment] = []

    for index, block in enumerate(_iter_block_items(document), start=1):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                heading = text.splitlines()[0][:255]
                segments.append(
                    ParsedSegment(
                        block_id=f"docx-fallback-{index}",
                        heading=heading,
                        content=text,
                        anchor={"page": 1, "section": heading, "line_start": index, "line_end": index},
                        block_type="paragraph",
                    )
                )
            continue

        segments.extend(_table_to_segments(block, start_index=len(segments) + 1))

    return ParsedDocument(parser_name="python_docx", parser_version="v3", segments=segments)


def _table_to_segments(table: Table, start_index: int) -> list[ParsedSegment]:
    segments: list[ParsedSegment] = []
    for offset, row in enumerate(table.rows):
        cells = [_cell_to_text(cell) for cell in row.cells]
        if any(cells):
            content = f"| {' | '.join(cells)} |"
            heading = cells[0][:255] if cells and cells[0] else f"table-row-{start_index + offset}"
            segments.append(
                ParsedSegment(
                    block_id=f"docx-table-{start_index + offset}",
                    heading=heading,
                    content=content,
                    anchor={"page": 1, "section": heading, "line_start": start_index + offset, "line_end": start_index + offset},
                    block_type="table_row",
                )
            )
    return segments


def _cell_to_text(cell: _Cell) -> str:
    parts: list[str] = []
    for block in _iter_block_items(cell):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                parts.append(text)
            continue

        parts.extend(segment.content for segment in _table_to_segments(block, start_index=1))
    return ' '.join(parts).strip()


def _iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_element = parent.element.body
    else:
        parent_element = parent._tc

    for child in parent_element.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def _segments_from_text(text: str, *, parser_name: str) -> list[ParsedSegment]:
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    segments: list[ParsedSegment] = []
    for index, block in enumerate(blocks, start=1):
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        heading = lines[0][:255]
        block_type = "table_row" if heading.startswith("|") else "paragraph"
        segments.append(
            ParsedSegment(
                block_id=f"{parser_name}-{index}",
                heading=heading,
                content=" ".join(lines) if block_type == "paragraph" else "\n".join(lines),
                anchor={"page": 1, "section": heading, "line_start": index, "line_end": index},
                block_type=block_type,
            )
        )
    return segments
