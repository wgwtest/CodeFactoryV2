from pathlib import Path

from docx import Document

from app.parsing.parsers import docx_parser


def _create_docx(path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value

    document.save(path)


def test_parse_docx_extracts_table_content_from_complex_document(tmp_path: Path) -> None:
    path = tmp_path / "complex.docx"
    _create_docx(
        path,
        paragraphs=["版本历史"],
        table_rows=[
            ["运行节点", "信息交换"],
            ["塔台", "航班状态"],
        ],
    )

    lines = docx_parser.parse_docx(str(path))
    joined = "\n".join(lines)

    assert "版本历史" in joined
    assert "运行节点" in joined
    assert "信息交换" in joined
    assert "塔台" in joined
    assert "航班状态" in joined


def test_parse_docx_prefers_docling_when_available(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "docling.docx"
    _create_docx(path, paragraphs=["旧解析内容"])

    class FakeDoclingDocument:
        def export_to_markdown(self) -> str:
            return "# OV-2\n\n| 运行节点 | 信息交换 |\n| --- | --- |\n| 塔台 | 航班状态 |"

    class FakeConversionResult:
        document = FakeDoclingDocument()

    class FakeConverter:
        def __init__(self) -> None:
            self.calls: list[str] = []

        def convert(self, file_path: str) -> FakeConversionResult:
            self.calls.append(file_path)
            return FakeConversionResult()

    fake_converter = FakeConverter()
    monkeypatch.setattr(docx_parser, "_build_docling_converter", lambda: fake_converter)

    lines = docx_parser.parse_docx(str(path))

    assert fake_converter.calls == [str(path)]
    assert "# OV-2" in "\n".join(lines)
    assert "| 塔台 | 航班状态 |" in "\n".join(lines)
